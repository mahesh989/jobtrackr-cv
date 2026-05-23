"""
Internal API consumed exclusively by JobTrackr's Next.js routes.

All endpoints require an HMAC-SHA256 signature in X-Signature, computed with
the shared JOBTRACKR_HMAC_SECRET. cv-backend has no other auth surface and
is not exposed to browsers.

Endpoints:
  - POST /internal/analyze                  — kicks off the pipeline (BackgroundTask)
  - POST /internal/extract-cv-text          — pypdf extraction for a Storage object
  - POST /internal/categorise-cv            — BYOK skill categorisation
  - POST /internal/extract-voice-fingerprint — voice fingerprint extraction
  - POST /internal/extract-stories          — story extraction from CV
  - POST /internal/match-stories            — deterministic story-to-JD ranking
  - POST /internal/scrape-jd               — JD scraping helper
  - POST /internal/research-company         — company research (Tavily + scrape + AI distill)
  - POST /internal/select-company-fact      — deterministic fact ranking (no AI)
  - POST /internal/generate-cover-letter        — single-call cover letter pipeline (BackgroundTask)
  - POST /internal/generate-opening-variants    — 3-4 P1 opener variants, synchronous
"""
from __future__ import annotations

import asyncio
import io
import logging

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, status

from app.config import get_settings
from app.database import get_supabase
from app.schemas.internal import (
    AnalyzeRequest,
    AnalyzeResponse,
    CategoriseCvRequest,
    CategoriseCvResponse,
    ExtractCvTextRequest,
    ExtractCvTextResponse,
    ExtractStoriesRequest,
    ExtractVoiceFingerprintRequest,
    ExtractVoiceFingerprintResponse,
    ScrapeJdRequest,
    ScrapeJdResponse,
)
from app.schemas.stories import (
    ExtractStoriesResponse,
    MatchStoriesRequest,
    MatchStoriesResponse,
    ScoredStory,
)
from app.security.hmac import verify_hmac
from app.services.ai.client import AIClientError, make_ai_client
from app.services.cv.skill_categoriser import categorise_cv_skills
from app.services.stories.story_extractor import extract_stories
from app.services.stories.story_matcher import score_stories
from app.services.voice.voice_fingerprint import extract_voice_fingerprint
from app.services.pipeline.orchestrator import run_analysis_pipeline
from app.services.scraping.jd_scraper import JDScrapeError, scrape_jd
from app.services.company.researcher import CompanyResearchError, research_company
from app.services.company.fact_selector import select_facts
from app.schemas.company import (
    ResearchCompanyRequest,
    ResearchCompanyResponse,
    CompanyResearch,
    SelectCompanyFactRequest,
    SelectCompanyFactResponse,
    RankedFact,
)
from app.schemas.cover_letter import (
    GenerateCoverLetterRequest,
    GenerateCoverLetterResponse,
    GenerateOpeningVariantsRequest,
    GenerateOpeningVariantsResponse,
    VoiceRewriteEmailRequest,
    VoiceRewriteEmailResponse,
)
from app.services.ai.prompts.cover_letter.voice_email import (
    VOICE_EMAIL_SYSTEM,
    VOICE_EMAIL_USER_TEMPLATE,
)
from app.services.cover_letter.generator import run_cover_letter_pipeline
from app.services.cover_letter.variants import generate_opening_variants

logger = logging.getLogger(__name__)

router = APIRouter(
    prefix="/internal",
    tags=["internal"],
    dependencies=[Depends(verify_hmac)],
)


# ── /internal/analyze ─────────────────────────────────────────────────────────

@router.post(
    "/analyze",
    response_model=AnalyzeResponse,
    status_code=status.HTTP_202_ACCEPTED,
)
async def analyze(
    body: AnalyzeRequest,
    background_tasks: BackgroundTasks,
) -> AnalyzeResponse:
    """
    Accept a pipeline trigger. Returns 202 immediately; the pipeline runs as a
    FastAPI BackgroundTask and writes step results to analysis_runs.{run_id}
    via Supabase service-role.
    """
    logger.info(
        "received run %s (user=%s provider=%s jd_len=%d cv_len=%d)",
        body.run_id, body.user_id, body.ai_provider,
        len(body.jd_text), len(body.cv_text),
    )
    background_tasks.add_task(run_analysis_pipeline, body)
    return AnalyzeResponse(run_id=body.run_id)


# ── /internal/extract-cv-text ────────────────────────────────────────────────

def _extract_pdf_text_sync(pdf_bytes: bytes) -> str:
    """
    Sync pypdf extraction — must run in a worker thread so the event loop
    doesn't block on large PDFs.

    Fixes the original cv-magic bug: synchronous pypdf inside `async def`.
    """
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(pdf_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)
    return "\n\n".join(pages).strip()


def _extract_docx_text_sync(docx_bytes: bytes) -> str:
    """Sync python-docx extraction. Run in a worker thread (see above)."""
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # Tables in CVs often hold contact lines + experience blocks — include them.
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n\n".join(paragraphs).strip()


@router.post("/extract-cv-text", response_model=ExtractCvTextResponse)
async def extract_cv_text(body: ExtractCvTextRequest) -> ExtractCvTextResponse:
    """Download a PDF from Supabase Storage and return its plain-text extraction."""
    settings = get_settings()
    bucket = settings.SUPABASE_CV_BUCKET

    # Path arrives in the form 'cvs/<user_id>/<cv_id>.pdf' — strip the bucket
    # prefix if it's been included by mistake.
    storage_key = body.storage_path
    prefix = f"{bucket}/"
    if storage_key.startswith(prefix):
        storage_key = storage_key[len(prefix):]

    # supabase-py is synchronous — wrap the download in a worker thread.
    def _download() -> bytes:
        return get_supabase().storage.from_(bucket).download(storage_key)

    try:
        file_bytes = await asyncio.to_thread(_download)
    except Exception as exc:
        logger.warning("extract-cv-text: download failed for %s: %s", storage_key, exc)
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Could not fetch CV file: {exc}",
        ) from exc

    # Dispatch by extension. The Storage bucket only allows PDF + DOCX
    # (enforced at migration 013), so a stray .doc/.txt would be rejected
    # at upload time — but we still defensively check here.
    lower = storage_key.lower()
    if lower.endswith(".pdf"):
        cv_text = await asyncio.to_thread(_extract_pdf_text_sync, file_bytes)
    elif lower.endswith(".docx"):
        cv_text = await asyncio.to_thread(_extract_docx_text_sync, file_bytes)
    else:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Unsupported file extension for {storage_key} (expected .pdf or .docx)",
        )

    word_count = len(cv_text.split())
    logger.info(
        "extract-cv-text: %s → %d chars, %d words",
        storage_key, len(cv_text), word_count,
    )
    return ExtractCvTextResponse(cv_text=cv_text, word_count=word_count)


# ── /internal/scrape-jd ──────────────────────────────────────────────────────

@router.post("/categorise-cv", response_model=CategoriseCvResponse)
async def categorise_cv(body: CategoriseCvRequest) -> CategoriseCvResponse:
    """
    BYOK skill categorisation. Returns three lists — technical / soft_skills /
    domain_knowledge — extracted from the provided CV text by the AI provider
    the user has connected. JobTrackr calls this once at CV upload time.
    """
    try:
        ai_client = make_ai_client(body.ai_provider, body.ai_api_key, body.ai_model)
    except AIClientError as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc))

    try:
        result = await categorise_cv_skills(ai_client, body.cv_text)
    except AIClientError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"AI categorisation failed: {exc}",
        )
    except ValueError as exc:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=str(exc))

    return CategoriseCvResponse(
        technical=        result.get("technical", []),
        soft_skills=      result.get("soft_skills", []),
        domain_knowledge= result.get("domain_knowledge", []),
    )


# ── /internal/extract-voice-fingerprint ──────────────────────────────────────

@router.post(
    "/extract-voice-fingerprint",
    response_model=ExtractVoiceFingerprintResponse,
)
async def extract_voice_fingerprint_endpoint(
    body: ExtractVoiceFingerprintRequest,
) -> ExtractVoiceFingerprintResponse:
    """
    Extract a structured voice fingerprint from a writing sample.

    Runs a deterministic trust score on the sample, then calls the user's
    AI provider (BYOK) to extract a 14-key fingerprint. Both the trust
    score and the fingerprint are returned; the caller (web API route) is
    responsible for persisting them to voice_profiles via service-role.

    NOTE: voice_sample_text must not appear in logs. If request-body logging
    is ever added to this service, add this field to the redaction list.
    """
    try:
        ai_client = make_ai_client(body.ai_provider, body.ai_api_key, body.ai_model)
    except AIClientError as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc))

    try:
        result = await extract_voice_fingerprint(ai_client, body.voice_sample_text)
    except AIClientError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"Voice fingerprint extraction failed: {exc}",
        )
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=str(exc)
        )

    return ExtractVoiceFingerprintResponse(
        fingerprint=result["fingerprint"],
        trust_score=result["trust_score"],
        trust_components=result["trust_components"],
        word_count=result["word_count"],
        matched_ai_phrases=result["matched_ai_phrases"],
    )


# ── /internal/extract-stories ────────────────────────────────────────────────

@router.post(
    "/extract-stories",
    response_model=ExtractStoriesResponse,
)
async def extract_stories_endpoint(
    body: ExtractStoriesRequest,
) -> ExtractStoriesResponse:
    """
    Extract structured achievement stories from a master CV.

    Calls the user's AI provider (BYOK) to identify 3–8 distinct achievements
    suitable for use as cover letter narratives. Validates each story against
    the Story Pydantic schema before returning. Returns HTTP 200 with an empty
    stories list and a diagnostic message if no achievements are found — this
    is not an error condition.

    NOTE: body.cv_text must not appear in logs. If request-body logging is
    ever added to this service, add cv_text to the redaction list.
    """
    try:
        ai_client = make_ai_client(body.ai_provider, body.ai_api_key, body.ai_model)
    except AIClientError as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc))

    try:
        result = await extract_stories(ai_client, body.cv_text)
    except AIClientError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"Story extraction failed: {exc}",
        )
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=str(exc)
        )

    return ExtractStoriesResponse(
        stories=result["stories"],
        diagnostic=result["diagnostic"],
    )


# ── /internal/match-stories ──────────────────────────────────────────────────

@router.post("/match-stories", response_model=MatchStoriesResponse)
async def match_stories_endpoint(body: MatchStoriesRequest) -> MatchStoriesResponse:
    """
    Rank stories against a JD using deterministic keyword overlap. No AI call.

    Caller (web route) passes the user's current story batch (with DB ids set)
    and the JD text. Returns scored story ids sorted by relevance descending.
    The web route merges scores back onto the full story objects by id.

    jd_text is treated as PII-adjacent (contains employer details) — only
    its length is logged, never the raw content.
    """
    logger.info(
        "match-stories: jd_len=%d stories=%d",
        len(body.jd_text),
        len(body.stories),
    )
    raw_scored = score_stories(
        body.jd_text,
        [s.model_dump() for s in body.stories],
    )
    return MatchStoriesResponse(
        scored=[
            ScoredStory(story_id=item["story_id"], score=item["score"])
            for item in raw_scored
        ]
    )


@router.post("/scrape-jd", response_model=ScrapeJdResponse)
async def scrape_jd_endpoint(body: ScrapeJdRequest) -> ScrapeJdResponse:
    """Scrape a job-posting URL for the cleaned JD text."""
    try:
        result = await scrape_jd(str(body.url))
    except JDScrapeError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=str(exc),
        ) from exc

    return ScrapeJdResponse(
        jd_text=result.jd_text,
        job_title=result.job_title,
        source_url=result.source_url,
    )


# ── /internal/research-company ────────────────────────────────────────────────

@router.post(
    "/research-company",
    response_model=ResearchCompanyResponse,
    status_code=status.HTTP_200_OK,
)
async def research_company_endpoint(body: ResearchCompanyRequest) -> ResearchCompanyResponse:
    """
    Research a company via Tavily + scraping + AI distillation.

    Returns the completed CompanyResearch synchronously (unlike /analyze
    which uses BackgroundTasks). Caller (web route) writes the result to
    Supabase and handles TTL/cache logic. Typical latency: 15–45 s.

    company_name length is logged; company_domain and ai_api_key are not.
    """
    settings = get_settings()

    logger.info(
        "research-company: company=%r domain_hint=%s provider=%s",
        body.company_name,
        bool(body.company_domain),
        body.ai_provider,
    )

    try:
        ai_client = make_ai_client(
            provider=body.ai_provider,
            api_key=body.ai_api_key,
            model=body.ai_model,
        )
    except AIClientError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Invalid AI client configuration: {exc}",
        ) from exc

    try:
        result_dict = await research_company(
            client=ai_client,
            company_name=body.company_name,
            company_domain=body.company_domain,
            tavily_api_key=settings.TAVILY_API_KEY or None,
        )
    except CompanyResearchError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=str(exc),
        ) from exc

    research = CompanyResearch(**result_dict)
    return ResearchCompanyResponse(
        company_id=research.company_id,
        status="completed",
        research=research,
        search_skipped=research.search_skipped,
    )


# ── /internal/select-company-fact ─────────────────────────────────────────────

@router.post(
    "/select-company-fact",
    response_model=SelectCompanyFactResponse,
    status_code=status.HTTP_200_OK,
)
async def select_company_fact_endpoint(body: SelectCompanyFactRequest) -> SelectCompanyFactResponse:
    """
    Rank company facts by keyword relevance to JD + CV. No AI call.

    Deterministic — same inputs always produce the same ranking.
    jd_text and cv_text lengths are logged; content is not.
    """
    logger.info(
        "select-company-fact: company_id=%r jd_len=%d cv_len=%d",
        body.company_id,
        len(body.jd_text),
        len(body.cv_text),
    )
    ranked = select_facts(
        jd_text=body.jd_text,
        cv_text=body.cv_text,
        facts=body.facts,
    )
    return SelectCompanyFactResponse(
        ranked_facts=[
            RankedFact(
                fact_text=item["fact_text"],
                score=item["score"],
                source_field=item["source_field"],
            )
            for item in ranked
        ]
    )


# ── /internal/generate-opening-variants ───────────────────────────────────────

@router.post(
    "/generate-opening-variants",
    response_model=GenerateOpeningVariantsResponse,
    status_code=status.HTTP_200_OK,
)
async def generate_opening_variants_endpoint(
    body: GenerateOpeningVariantsRequest,
) -> GenerateOpeningVariantsResponse:
    """
    Generate 3-4 structurally distinct P1 openers in a single AI call.

    Unlike /generate-cover-letter this endpoint is synchronous — it returns
    the variants in the response body (typical latency: 5-15 s). The caller
    (web /cover-letter POST route) stores the variants in the cover_letters
    row and returns them to the browser for the picker UI.

    NOTE: body.voice_sample_text must not appear in logs.
    """
    logger.info(
        "generate-opening-variants: user=%s job=%s provider=%s",
        body.user_id, body.job_id, body.ai_provider,
    )

    try:
        ai_client = make_ai_client(body.ai_provider, body.ai_api_key, body.ai_model)
    except AIClientError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Invalid AI client configuration: {exc}",
        ) from exc

    try:
        variants = await generate_opening_variants(ai_client, body)
    except AIClientError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"Opening variants generation failed: {exc}",
        ) from exc

    return GenerateOpeningVariantsResponse(variants=variants)


# ── /internal/generate-cover-letter ───────────────────────────────────────────

@router.post(
    "/generate-cover-letter",
    response_model=GenerateCoverLetterResponse,
    status_code=status.HTTP_202_ACCEPTED,
)
async def generate_cover_letter(
    body: GenerateCoverLetterRequest,
    background_tasks: BackgroundTasks,
) -> GenerateCoverLetterResponse:
    """
    Accept a cover letter generation trigger. Returns 202 immediately.

    The three-pass pipeline (skeleton → voice transfer → burstiness) runs as a
    FastAPI BackgroundTask and writes progress + outputs to cover_letters.{letter_id}
    via Supabase service-role. The browser subscribes to postgres_changes on
    cover_letters for real-time progress (same pattern as analysis_runs).

    NOTE: body.voice_sample_text must not appear in logs. See GenerateCoverLetterRequest
    privacy annotation.
    """
    logger.info(
        "generate-cover-letter: letter_id=%s user=%s provider=%s jd_len=%d",
        body.letter_id, body.user_id, body.ai_provider, len(body.jd_text),
    )
    background_tasks.add_task(run_cover_letter_pipeline, body)
    return GenerateCoverLetterResponse(letter_id=body.letter_id)


# ── /internal/voice-rewrite-email ─────────────────────────────────────────────

@router.post(
    "/voice-rewrite-email",
    response_model=VoiceRewriteEmailResponse,
    status_code=status.HTTP_200_OK,
)
async def voice_rewrite_email_endpoint(
    body: VoiceRewriteEmailRequest,
) -> VoiceRewriteEmailResponse:
    """
    Rewrite the SHORT email cover note that ships an application, in the
    candidate's voice. Synchronous — one AI call, returns the body text.

    The web tier calls this from /api/applications/[letter_id]/email-draft
    when the cached email_body is null and a voice_sample_raw exists. The
    result is cached in cover_letters.email_body so subsequent draft loads
    are instant.

    PRIVACY: body.voice_sample_text must not appear in logs.
    """
    logger.info(
        "voice-rewrite-email: user=%s letter=%s provider=%s job_title=%r company=%r",
        body.user_id, body.letter_id, body.ai_provider, body.job_title, body.company,
    )

    try:
        ai_client = make_ai_client(body.ai_provider, body.ai_api_key, body.ai_model)
    except AIClientError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Invalid AI client configuration: {exc}",
        ) from exc

    user_prompt = VOICE_EMAIL_USER_TEMPLATE.format(
        voice_sample=body.voice_sample_text,
        boilerplate=body.boilerplate_body,
    )

    try:
        rewritten = await ai_client.complete(
            system=VOICE_EMAIL_SYSTEM,
            user=user_prompt,
            max_tokens=800,
            # Style transfer is more constrained than free-form generation —
            # we want the same meaning every time, just reshaped. Lower temp
            # also makes the AI less likely to drift into autobiography.
            temperature=0.3,
            no_training=True,
        )
    except AIClientError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"Voice rewrite failed: {exc}",
        ) from exc

    cleaned = rewritten.strip()
    if not cleaned:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail="Voice rewrite returned empty body",
        )

    return VoiceRewriteEmailResponse(body=cleaned)
