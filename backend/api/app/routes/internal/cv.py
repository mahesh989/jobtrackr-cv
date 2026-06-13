from __future__ import annotations

import asyncio
import io
import logging

from fastapi import APIRouter, HTTPException, status

from app.config import get_settings
from app.database import get_supabase
from app.schemas.internal import (
    CategoriseCvRequest,
    CategoriseCvResponse,
    CvReferee,
    ExtractCvReferencesRequest,
    ExtractCvReferencesResponse,
    ExtractCvTextRequest,
    ExtractCvTextResponse,
)
from app.services.ai.client import AIClientError, make_ai_client
from app.services.cv.skill_categoriser import categorise_cv_skills
from app.services.cv.references_extractor import extract_cv_references

logger = logging.getLogger(__name__)

router = APIRouter()

def _extract_pdf_text_sync(pdf_bytes: bytes) -> str:
    """
    Sync pypdf extraction — must run in a worker thread so the event loop
    doesn't block on large PDFs.

    Fixes the original cv-magic bug: synchronous pypdf inside `async def`.
    """
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(pdf_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)
    return "\n\n".join(pages).strip()


def _extract_docx_text_sync(docx_bytes: bytes) -> str:
    """Sync python-docx extraction. Run in a worker thread (see above)."""
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # Tables in CVs often hold contact lines + experience blocks — include them.
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n\n".join(paragraphs).strip()


@router.post("/extract-cv-text", response_model=ExtractCvTextResponse)
async def extract_cv_text(body: ExtractCvTextRequest) -> ExtractCvTextResponse:
    """Download a PDF from Supabase Storage and return its plain-text extraction."""
    settings = get_settings()
    bucket = settings.SUPABASE_CV_BUCKET

    # Path arrives in the form 'cvs/<user_id>/<cv_id>.pdf' — strip the bucket
    # prefix if it's been included by mistake.
    storage_key = body.storage_path
    prefix = f"{bucket}/"
    if storage_key.startswith(prefix):
        storage_key = storage_key[len(prefix):]

    # supabase-py is synchronous — wrap the download in a worker thread.
    def _download() -> bytes:
        return get_supabase().storage.from_(bucket).download(storage_key)

    try:
        file_bytes = await asyncio.to_thread(_download)
    except Exception as exc:
        logger.warning("extract-cv-text: download failed for %s: %s", storage_key, exc)
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Could not fetch CV file: {exc}",
        ) from exc

    # Size cap — the signed-upload flow doesn't enforce one at the edge, so a
    # huge/crafted file could otherwise be handed straight to pypdf/python-docx.
    _MAX_CV_BYTES = 10 * 1024 * 1024  # 10 MB — generous; real CVs are ~80-300 KB
    if len(file_bytes) > _MAX_CV_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="CV file is too large.",
        )

    # Dispatch by extension AND verify the content actually matches that type
    # by its magic bytes. The extension alone is attacker-controllable (the
    # browser PUTs bytes directly to Storage via the signed URL), so a non-PDF
    # payload could arrive as .pdf. PDF → "%PDF"; DOCX is a ZIP → "PK\x03\x04".
    lower = storage_key.lower()
    if lower.endswith(".pdf"):
        if not file_bytes.startswith(b"%PDF"):
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail="File is not a valid PDF.",
            )
        cv_text = await asyncio.to_thread(_extract_pdf_text_sync, file_bytes)
    elif lower.endswith(".docx"):
        if not file_bytes.startswith(b"PK\x03\x04"):
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail="File is not a valid DOCX.",
            )
        cv_text = await asyncio.to_thread(_extract_docx_text_sync, file_bytes)
    else:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Unsupported file extension for {storage_key} (expected .pdf or .docx)",
        )

    word_count = len(cv_text.split())
    logger.info(
        "extract-cv-text: %s → %d chars, %d words",
        storage_key, len(cv_text), word_count,
    )
    return ExtractCvTextResponse(cv_text=cv_text, word_count=word_count)


# ── /internal/scrape-jd ──────────────────────────────────────────────────────

@router.post("/categorise-cv", response_model=CategoriseCvResponse)
async def categorise_cv(body: CategoriseCvRequest) -> CategoriseCvResponse:
    """
    BYOK skill categorisation. Returns three lists — technical / soft_skills /
    domain_knowledge — extracted from the provided CV text by the AI provider
    the user has connected. JobTrackr calls this once at CV upload time.
    """
    try:
        ai_client = make_ai_client(body.ai_provider, body.ai_api_key, body.ai_model)
    except AIClientError as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc))

    try:
        result = await categorise_cv_skills(ai_client, body.cv_text)
    except AIClientError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"AI categorisation failed: {exc}",
        )
    except ValueError as exc:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=str(exc))

    return CategoriseCvResponse(
        technical=        result.get("technical", []),
        soft_skills=      result.get("soft_skills", []),
        domain_knowledge= result.get("domain_knowledge", []),
    )


# ── /internal/extract-cv-references ──────────────────────────────────────────

@router.post("/extract-cv-references", response_model=ExtractCvReferencesResponse)
async def extract_cv_references_route(
    body: ExtractCvReferencesRequest,
) -> ExtractCvReferencesResponse:
    """
    BYOK referee extraction from CV text. Returns up to 3 referees with
    {name, job_title, company, email}. Called on-demand from the web UI
    when a user clicks "Extract from active CV" in the References section.
    """
    try:
        ai_client = make_ai_client(body.ai_provider, body.ai_api_key, body.ai_model)
    except AIClientError as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc))

    try:
        referees = await extract_cv_references(ai_client, body.cv_text)
    except AIClientError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"AI extraction failed: {exc}",
        )
    except ValueError as exc:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=str(exc))

    return ExtractCvReferencesResponse(
        referees=[CvReferee(**r) for r in referees],
    )


# ── /internal/extract-voice-fingerprint ──────────────────────────────────────


