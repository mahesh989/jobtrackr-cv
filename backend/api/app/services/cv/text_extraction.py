"""CV file → plain-text extraction (PDF / DOCX) + content validation.

Extracted verbatim from routes/internal/cv.py so the route stays a thin
transport layer. All extractors are SYNC — callers must wrap in
asyncio.to_thread so the event loop doesn't block on large files.
"""
from __future__ import annotations

import io

# Magic-byte signatures — the extension alone is attacker-controllable (the
# browser PUTs bytes directly to Storage via the signed URL), so a non-PDF
# payload could arrive as .pdf. PDF → "%PDF"; DOCX is a ZIP → "PK\x03\x04".
PDF_MAGIC = b"%PDF"
DOCX_MAGIC = b"PK\x03\x04"


def extract_pdf_text_sync(pdf_bytes: bytes) -> str:
    """
    Sync pypdf extraction — must run in a worker thread so the event loop
    doesn't block on large PDFs.

    Fixes the original cv-magic bug: synchronous pypdf inside `async def`.
    """
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(pdf_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)
    return "\n\n".join(pages).strip()


def extract_docx_text_sync(docx_bytes: bytes) -> str:
    """Sync python-docx extraction. Run in a worker thread (see above)."""
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # Tables in CVs often hold contact lines + experience blocks — include them.
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n\n".join(paragraphs).strip()
